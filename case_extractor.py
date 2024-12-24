import os
import asyncio
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import List, Tuple, Optional
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
import dotenv
from functools import partial

from openai import OpenAI

import instructor
from pydantic import BaseModel, Field
from collections import defaultdict

dotenv.load_dotenv()

# Configuration
INPUT_FOLDER = "input_docs"
OUTPUT_FOLDER = "analysis_results"

MAX_CONCURRENT_REQUESTS = int(os.getenv("MAX_CONCURRENT_REQUESTS", "10"))
MAX_RETRIES = int(os.getenv("MAX_RETRIES", "1"))
RETRY_DELAY = int(os.getenv("RETRY_DELAY", "1"))

# OpenAI model name can be specified if needed
# For example, you might set an environment variable OPENAI_MODEL_NAME, or just hardcode a model name.
OPENAI_MODEL_NAME = os.getenv("OPENAI_MODEL_NAME", "o1")

# Thread pool for concurrency
thread_pool = ThreadPoolExecutor(max_workers=MAX_CONCURRENT_REQUESTS)

class ExtractCaseRelevancy(BaseModel):
    blue_book_citation: str = Field(..., description="The full Blue Book style citation for this case.")
    summary: str = Field(..., description="A brief summary of the case, 2-4 sentences.")
    relevance_level: str = Field(..., description="One of 'High', 'Medium', or 'Low' relevancy to the issue of whether intentional spoliation should be recognized as a standalone tort after Reynolds v. Bordelon.")
    reasoning: str = Field(..., description="Explanation of why it was assigned this relevance level.")
    key_points: List[str] = Field(..., description="A list of key points or mentions in the case related to spoliation.")
    citations: List[str] = Field(..., description="List of the key cases cited in the opinion.")
    quotes: List[str] = Field(..., description="Key quotes (if any) from the case that support or refute the argument about intentional spoliation not being a standalone tort, cited in blue book style to the exact page number.")
    argument: str = Field(..., description="A concise argument explaining how this case supports the Reynolds court's rationale that intentional spoliation should not be recognized as a standalone tort, or if the case does not support the argument, attempt to distinguish the case. IMPORTANT:The Argument section should be styled like the analysis section of a legal brief. Every citation to a fact or conclusion of law must be supported by a Blue Book citation to the case being referenced.")
    support_level: str = Field(..., description="One of 'Strongly Supports', 'Supports', or 'Does not Support'.")

class SayHi(BaseModel):
    hi: str = Field(..., description="Say hi")

@dataclass
class DocumentAnalysis:
    filename: str
    analysis: Optional[ExtractCaseRelevancy]
    error: Optional[str] = None

def extract_text_from_docx(file_path):
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return text

def create_formatted_docx(doc, filename, analysis: ExtractCaseRelevancy, is_first=False):
    """Add a formatted analysis to the Word document."""
    if not is_first:
        # Add page break between documents
        doc.add_page_break()
    
    # Add title
    title = doc.add_heading(f'Case Analysis: {analysis.blue_book_citation}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add horizontal line
    doc.add_paragraph('_' * 50)
    
    # SUMMARY
    heading = doc.add_heading('SUMMARY', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    summary_para = doc.add_paragraph(analysis.summary)
    summary_para.paragraph_format.first_line_indent = Inches(0.25)

    # RELEVANCY
    heading = doc.add_heading('RELEVANCY', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    rel_para = doc.add_paragraph()
    rel_para.add_run(f"Relevance Level: {analysis.relevance_level}\n")
    rel_para.add_run(f"Reasoning: {analysis.reasoning}\n")

    # SUPPORT LEVEL
    heading = doc.add_heading('SUPPORT LEVEL', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    supp_para = doc.add_paragraph()
    supp_para.add_run(f"Support Level: {analysis.support_level}\n")

    # ARGUMENT
    heading = doc.add_heading('ARGUMENT', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    arg_para = doc.add_paragraph(analysis.argument)
    arg_para.paragraph_format.first_line_indent = Inches(0.25)

    # KEY POINTS
    heading = doc.add_heading('KEY POINTS RELATED TO SPOLIATION', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    for point in analysis.key_points:
        kp = doc.add_paragraph(point, style='List Bullet')
        kp.paragraph_format.left_indent = Inches(0.5)

    # CITATIONS
    heading = doc.add_heading('CITATIONS', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    for cite in analysis.citations:
        cp = doc.add_paragraph(cite, style='List Bullet')
        cp.paragraph_format.left_indent = Inches(0.5)

    # QUOTES
    heading = doc.add_heading('QUOTES', level=2)
    heading.style.font.color.rgb = RGBColor(0, 51, 102)
    for quote in analysis.quotes:
        qp = doc.add_paragraph(quote, style='List Bullet')
        qp.paragraph_format.left_indent = Inches(0.5)

    return doc

async def analyze_text_with_instructor(client, text: str, filename: str, retry_count: int = 0) -> DocumentAnalysis:
    """
    Analyze text using an OpenAI model via instructor.
    """
    print(f"Retry count: {retry_count}")

    system_prompt = """You are a legal research assistant with extensive knowledge of Louisiana tort law and the reasoning in Reynolds v. Bordelon. 

You will be given the text of a Louisiana court case that uses either the phrase "intentional spoliation" or "impairment of a civil claim."

As you know, after the Louisiana Supreme Court decided Reynolds v. Bordelon in 2015, Louisiana does not recognize an independent tort for negligent spoliation. We need research done so that we can argue that the same reasoning should apply to intentional spoliation also, meaning neither intentional spoliation nor negligent spoliation should be considered a standalone tort. We are trying to find case law that supports this argument. The cases that contain the phrase "impairment of a civil claim" are also relevant to this research because historically spoliation claims were sometimes referred to by this phrase. 

Your task:
1. Extract the full Blue Book style citation for this case.

2. Determine how relevant this case is to supporting the argument that after Reynolds v. Bordelon, intentional spoliation should not be recognized as a standalone tort in Louisiana. IMPORTANT: A case can be relevant if it does *not* support our arguement. For example, if a court finds that intentional spoliation is a standalone tort, that is highly relevant, but it should also be marked as "Does not support."

3. Rate relevance as High, Medium, or Low:

   - High: The case directly addresses whether intentional spoliation should be recognized as a standalone tort, regardless of whether it supports or does not support our argument. 

   - Medium: The case references spoliation and tort claims but does not directly confirm or refute the standalone tort nature for intentional spoliation after Reynolds.

   - Low: The case barely touches on the nature of spoliation as a tort or is tangential and does not aid the argument.

4. Provide a brief summary, a reasoning for the chosen relevance level, key points related to spoliation found in the text, any relevant citations, and any direct quotes that could help build the argument.

5. Bear in mind that some of these cases (if not most) likely pre-date Reynolds v. Bordelon, so the idea is to figure out how an earlier case can be said to support Reynolds even if it predates that decision.

6. Craft a concise argument that explains how this case's reasoning, holdings, or dicta support the Reynolds court's rationale that intentional spoliation should not be recognized as a standalone tort. Focus on drawing parallels between the case's treatment of spoliation or impairment of a civil claim and the Reynolds court's concerns about duplicative recovery, existing remedies, and the policy implications of recognizing new causes of action. If the case truly cannot be argued to support the argument, then indicate that in the argument section and try to distinguish the case. IMPORTANT:The Argument section should be styled like the analysis section of a legal brief. Every citation to a fact or conclusion of law must be supported by a Blue Book citation to the case being referenced.

7. Finally, indicate one of the following for "support_level": 'Strongly Supports', 'Supports', or 'Does not Support' to reflect how strongly this case supports the argument about intentional spoliation not being recognized as a standalone tort after Reynolds.

Follow the structure and fields defined in the structured output model:
blue_book_citation, summary, relevance_level, reasoning, key_points, citations, quotes, argument, support_level.
"""

    user_prompt = f"""
    # REYNOLDS v. BORDELAN CASE FOR REFERENCE:

    **1 The instant case presents a claim under the Louisiana Products Liability Act (“LPLA”). We granted its companion case to determine the viability of negligent spoliation of evidence as a cause of action in Louisiana.1 We now address the underlying products liability case and review the appropriateness of the lower court's grant of summary judgment. For the reasons expressed below, we affirm.

FACTS AND PROCEDURAL HISTORY
On March 15, 2008, a multi-vehicle accident occurred in St. Tammany **2 Parish.
 *610 Robert J. Bordelon, III is alleged to have caused the accident when he swerved two separate times from the left lane of traffic to the right lane, colliding with two vehicles. The second collision involved the instant plaintiff, Richard Reynolds, who was driving a 2003 Infiniti G35S, which was manufactured by Nissan North America (“Nissan”). After the initial impact, the plaintiff was pushed into another vehicle and came to rest in a ditch. The accident caused the plaintiff to sustain serious injuries.
On March 12, 2009, the plaintiff filed suit against Bordelon and other defendants. With regard to Nissan, the plaintiff asserted claims under the LPLA for the failure of the air bags to deploy and/or operate. Specifically, he alleged the Infiniti was defective (1) due to a construction or composition defect; (2) due to a design defect; (3) for failure to contain an adequate warning; and (4) for failure to conform to an express warranty.
On July 8, 2013, Nissan filed a motion for summary judgment. In opposition, the plaintiff filed the affidavit of Dr. Richard Baratta. Ultimately, the trial court made several evidentiary rulings and concluded that there was an absence of factual support for any of the product liability theories, and it granted summary judgment in favor of Nissan. The court of appeal affirmed the judgment, finding no error in the evidentiary rulings and that there were no genuine issues of material fact upon which to survive summary judgment.2 We granted the plaintiff's writ application to review the grant of summary judgment.3
APPLICABLE LAW
1
2
A motion for summary judgment is a procedural device used when there is no genuine issue of material fact for all or part of the relief prayed for by a **3 litigant.4 A summary judgment is reviewed on appeal de novo, with the appellate court using the same criteria that govern the trial court's determination of whether summary judgment is appropriate; i.e. whether there is any genuine issue of material fact, and whether the movant is entitled to judgment as a matter of law.5
3
A motion for summary judgment will be granted “if the pleadings, depositions, answers to interrogatories, and admissions on file, together with the affidavits, if any, show that there is no genuine issue as to material fact, and that mover is entitled to judgment as a matter of law.”6 The burden of proof remains with the movant. However, if the movant will not bear the burden of proof at trial on the matter that is before the court on the motion for summary judgment, the movant's burden on the motion does not require him to negate all essential elements of the adverse party's claim, action, or defense, but rather to point out to the court that there is an absence of factual support for one or more elements essential to the adverse party's claim, action, or defense. Thereafter, if the adverse party fails to produce factual support sufficient to establish that he will be able to satisfy his evidentiary *611 burden of proof at trial, there is no genuine issue of material fact.7
This court explained the summary judgment procedure as follows:
[The summary judgment procedure] first places the burden of producing evidence at the hearing on the motion for summary judgment on the mover (normally the defendant), who can ordinarily meet that burden by submitting affidavits or by pointing out the lack of factual support for an essential element in the opponent's case. At that point, the party who bears the burden of persuasion at trial (usually the plaintiff) must come forth with evidence (affidavits or discovery responses) which demonstrates he or she will be able to **4 meet the burden at trial.... Once the motion for summary judgment has been properly supported by the moving party, the failure of the non-moving party to produce evidence of a material factual dispute mandates the granting of the motion. (Emphasis added; citation omitted).8
DISCUSSION
Evidentiary Rulings
The plaintiff avers that the trial court improperly excluded evidence in its consideration of the motion for summary judgment. We will address each category of evidence that the lower courts ruled inadmissible.
First, the plaintiff attempted to admit post-accident photographs of the vehicle. However, the photographs were not verified or authenticated in any way. The plaintiff did not introduce an affidavit or testimony by any person familiar with the photographs, the photographer or otherwise, in order to lay the foundation that the photographs were actually depictions of the plaintiff's vehicle or that the depictions were accurate.9 Accordingly, we agree that the lack of verification deems the photographs inadmissible at the summary judgment hearing.
Next, the plaintiff argues the airbag service bulletin printed from the National Highway Traffic Safety Administration (NHTSA) should have been admissible. However, no corroborating testimony or affidavit was presented to establish the printout's authenticity. As noted by the court of appeal, the front page of the printout states that public documents were unavailable at the time the printout was produced; thus, the plaintiff cannot avail himself of any self-authenticating public document exception to survive this hurdle of admissibility.10 Accordingly, we find no error in the evidentiary ruling that excluded the NHTSA service bulletin.
**5 Further, the plaintiff sought to admit other incident investigation reports created by Nissan, presumably in an effort to illustrate that Nissan knew of an ongoing problem of its vehicles' airbags failing to deploy and failed to warn consumers. However, as articulated by the court of appeal, the plaintiff failed to establish the reports' relevancy to this proceeding. The reports do not reference the plaintiff's accident, nor do they demonstrate any similarities to the plaintiff's specific vehicle or the instant circumstances surrounding the alleged failure of the airbags to deploy. Rather, the investigation reports pertain to varying makes and models of vehicles in varying years in varying locales. Thus, we *612 find the reports are not relevant evidence as they do not have a tendency to make the existence of a material fact more probable or less probable than such a determination would be without the evidence.11
The plaintiff also attempted to introduce pre-accident invoices for service performed on his vehicle by the car dealership that sold him his vehicle, Ray Brandt Infiniti. While the documents purportedly are records made and kept in the course of regularly conducted business activity for purposes of the business records hearsay exception, the plaintiff introduced no affidavit of the custodian or any other witness to corroborate their nature.12 Accordingly, the invoices were properly excluded.
Additionally, the plaintiff, in opposing the motion for summary judgment, sought to introduce the affidavit and the curriculum vitae of his expert in accident reconstruction, Dr. Richard V. Baratta, Ph.D., P.E. The trial court admitted into the evidence the affidavit, in which Dr. Baratta opines the “airbags should have deployed to assist in mitigating [the plaintiff's] injuries.” However, the trial court **6 found his curriculum vitae inadmissible because it was unsworn and uncertified.13 Our own review of the curriculum vitae supports the finding that it is not a sworn or certified copy and is, thus, inadmissible. Additionally, the materials reviewed by Dr. Baratta, were unsworn and uncertified and bore the added defect of not being referenced in the curriculum vitae. Accordingly, they were properly ruled inadmissible as well.
Last, we find evidentiary problems with emailed materials sent by Dr. Baratta to the plaintiff's attorney. Allegedly, the materials were used to support the conclusions drawn in Dr. Baratta's affidavit; however, they, too, are not authenticated, irrelevant, and constitute hearsay. Accordingly, the trial court properly excluded these emails.
Review of the Merits
Having established what evidence is and is not before us, we turn now to the motion for summary judgment. Nissan, as the movant, must satisfy his burden by “submitting affidavits or pointing out the lack of factual support for an essential element in the opponent's case.”14 The plaintiff, as the party who bears the burden of proof at trial, must then “come forth with evidence which demonstrates he will be able to meet the burden at trial.”
The plaintiff's case arises under the LPLA, which provides the exclusive theories under which a plaintiff can pursue a claim against a manufacturer for an alleged product defect.15 La.R. 9:2800.54 sets forth the elements of a products liability claim:
A. The manufacturer of a product shall be liable to a claimant for damage proximately caused by a characteristic of the product that renders the product unreasonably dangerous when such damage arose from a reasonably anticipated use of the product by the **7 claimant or another person or entity.
B. A product is unreasonably dangerous if and only if:
(1) The product is unreasonably dangerous in construction or composition as provided in R.S. 9:2800.55;
*613 (2) The product is unreasonably dangerous in design as provided in R.S. 9:2800.56;
(3) The product is unreasonably dangerous because an adequate warning about the product has not been provided as provided in R.S. 9:2800.57; or
(4) The product is unreasonably dangerous because it does not conform to an express warranty of the manufacturer about the product as provided in R.S. 9:2800.58.
C. The characteristic of the product that renders it unreasonably dangerous under R.S. 9:2800.55 must exist at the time the product left the control of its manufacturer. The characteristic of the product that renders it unreasonably dangerous under R.S. 9:2800.56 or 9:2800.57 must exist at the time the product left the control of its manufacturer or result from a reasonably anticipated alteration or modification of the product.
D. The claimant has the burden of proving the elements of Subsections A, B and C of this Section.
4
Moving to the plaintiff's first alleged theory, a product is unreasonably dangerous in construction or composition if, at the time the product left its manufacturer's control, the product deviated in a material way from the manufacturer's specifications or performance standards for the product or from otherwise identical products manufactured by the same manufacturer.16 To prove this theory, the plaintiff must show (1) what Nissan's specifications or performance standards were for the vehicle/airbags and (2) how the plaintiff's vehicle/airbags materially deviated from these standards so as to render it unreasonably dangerous.17 In an effort to meet this burden of proof, the plaintiff **8 introduced the owner's manual for his vehicle. Relative to the supplemental restraint system/air bags, the manual provides:
These [air bag] systems are designed to meet voluntary guidelines to help reduce the risk of injury to out-of-position occupants ... The supplemental side air bags and curtain side-impact air bags are designed to inflate in higher severity side collisions on the side of the vehicle impacted, although they may inflate if the forces in another type of collision are similar to those of a higher severity side impact. They are designed to inflate on the side where the vehicle is impacted. They may not inflate in certain side collisions. Vehicle damage [or lack of it] is not always an indication of proper supplemental side air bag and curtain side-impact air bag operation.
To the extent the plaintiff is arguing the owner's manual qualifies as Nissans' performance standards, we note that, by its own acknowledgement, the air bags “may not inflate” in certain side collisions. This caveat necessarily precludes a finding that there existed a stated expectation of a specified performance.
Moreover, the manual's language that the air bags “are designed to inflate in higher severity side collisions on the side of the vehicle impacted” begs the question of whether the plaintiff's collision is a collision that falls within the expectations that would trigger inflation of the air bags. We find the plaintiff offered nothing to answer this inquiry in the affirmative so as to satisfy his burden of proof on a construction or composition defect. Dr. Baratta's affidavit makes a conclusory statement, unaided by factual support, that the *614 air bag should have deployed. This conclusion is insufficient on two grounds. First, the curriculum vitae that sought to establish Dr. Baratta's expertise, was ruled inadmissible. Thus, we find nothing to confirm his expertise as an airbag expert or even an accident re-constructionist. Second, Dr. Baratta's conclusory statements are not supported by any factual evidence and contain no explanation as to how he reached his conclusion. The list of reviewed materials, which perhaps could have established this factual requisite, also suffered the fate of inadmissibility. Accordingly, we find the plaintiff presented no evidence sufficient to establish he **9 would satisfy his evidentiary burden of proving a composition or construction defect.
5
The next theory under the LPLA advanced by the plaintiff is a design defect.18 A product is unreasonably dangerous in design if, at the time the product left its manufacturer's control:
(1) There existed an alternative design for the product that was capable of preventing the claimant's damage; and
(2) The likelihood that the product's design would cause the claimant's damage and the gravity of that damage outweighed the burden on the manufacturer of adopting such alternative design and the adverse effect, if any, of such alternative design on the utility of the product.19
Accordingly, the plaintiff was first required to show an alternative design for the supplemental restraint system existed at the time it left Nissan's control. The plaintiff proposed no other design for the product, and, indeed, admitted that he did not develop an alternative design. Without proving this required element, it is unnecessary to address the remaining elements of this theory insofar as the LPLA requires all elements to be proven.20
6
The plaintiff next alleged his vehicle and/or air bag system was unreasonably dangerous due to an inadequate warning. La.R.S. 9:2800.57(A) provides:
A product is unreasonably dangerous because an adequate warning about the product has not been provided if, at the time the product left its manufacturer's control, the product possessed a characteristic that may cause damage and the manufacturer failed to use reasonable care to provide an adequate warning of such characteristic and its danger to users and handlers of the product.
“Adequate warning” is defined as “a warning or instruction that would lead an ordinary reasonable user or handler of a product to contemplate the danger in using or handling the product and either to decline to use or handle the product or, if **10 possible, to use or handle the product in such a manner as to avoid the damage for which the claim is made.”21 The plaintiff did not specify what warning was inadequate, did not provide a proposed adequate warning, and did not provide any evidence to support this claim. The plaintiff stated in a discovery response that Nissan failed to provide an adequate warning that there existed a defect with the air bag system. However, this conclusory statement again fails for presuming a defect has been proven *615 when no evidence was presented to show the nature or existence of that defect.22
Inasmuch as the owner's manual contained a warning about the side air bag “ordinarily not inflat[ing] in the event of a frontal impact, rear impact, rollover or lower severity side collision,” we note the initial impact, as described in the police investigation report, was a rear collision, followed by a frontal collision. Thus, we find the warning contained in the owner's manual adequately explained that the side air bags typically do not deploy in frontal or rear collisions, precluding a finding of a genuine issue of material fact on this ground.
7
For each of the above theories, we also note, pursuant to La.R.S. 9:2800.54(C), the plaintiff must show that the vehicle/supplemental restraint system was in substantially the same condition as it was when it left its manufacturer's control and that there were no alterations or modifications to the vehicle's air bag system. Again, there is an absence of factual proof to assist the plaintiff in meeting this required evidentiary burden. The plaintiff's pleadings, responses to discovery requests, subsequent appellate briefs, and exhibits (including Dr. Baratta's affidavit) are silent as to the vehicle's condition at the time it left Nissan's control and are silent as to whether the air bag system was in an unaltered and unmodified condition at the time of the accident. The plaintiff **11 produces no evidence to show the alleged defect in the air bag system was not caused by the collision itself. Accordingly, we find the plaintiff failed to produce evidence that would demonstrate he could carry the burden of proof on these required elements, in addition to those articulated in the above-specified theories.
8
Last, the plaintiff brought a claim under the express warranty provision of the LPLA. “A product is unreasonably dangerous when it does not conform to an express warranty made at any time by the manufacturer about the product if the express warranty has induced the claimant or another person or entity to use the product and the claimant's damage was proximately caused because the express warranty was untrue.”23 The plaintiff did not identify (1) a specific express warranty that induced him to use his vehicle, (2) did not prove that the warranty was untrue, and (3) did not show that the failure to conform to that express warranty caused his injuries. Rather, the plaintiff alleged that he, as a consumer, bought the vehicle because Nissan warranted that “[the vehicle] will perform like it is supposed to.” As stated by the court of appeal, the plaintiff did not “point to any specific express warranty given by Nissan, but instead claim[ed] generally that Nissan had given him the expectation that his vehicle's air bag system would mitigate his injuries in a severe automobile accident, and it did not.”24 We cannot accept a general alleged warranty for purposes of an express warranty claim. The LPLA makes it very clear that in order for the manufacturer to be liable, there must be a specified stated warranty, i.e., express.
Dr. Baratta's affidavit opines that the owner's manual gave an expectation that in a high severity side impact, the side curtain air bags would deploy. However, and as discussed above, the manual specifically provided that the side air bags “may not *616 inflate in certain side collisions.” Thus, in the absence of an express **12 statement warranting to the plaintiff that his air bags would have deployed in a collision substantially similar to his own, we find he cannot prevail on this claim at trial.
CONCLUSION
For the reasons expressed herein, we find no error in the grant of summary judgment in favor of Nissan and we affirm the judgment of the court of appeal. AFFIRMED.

# THE CASE TEXT YOU ARE ANALYZING:

IMPORTANT:The Argument section should be styled like the analysis section of a legal brief. Every citation to a fact or conclusion of law must be supported by a Blue Book citation to the case being referenced.

{text}
"""

    for attempt in range(retry_count, MAX_RETRIES + 1):
        try:
            messages = [
                {"role": "developer", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]

            if OPENAI_MODEL_NAME == "o1":
                resp = await asyncio.get_event_loop().run_in_executor(
                    None,
                    partial(
                        client.messages.create,
                        messages=messages,
                        model=OPENAI_MODEL_NAME,
                        reasoning_effort="high",
                        response_model=ExtractCaseRelevancy,
                    )
                )
            else:
                resp = await asyncio.get_event_loop().run_in_executor(
                    None,
                    partial(
                        client.messages.create,
                        messages=messages,
                        model=OPENAI_MODEL_NAME,
                        response_model=ExtractCaseRelevancy,
                        temperature=0
                    )
                )
            print("✓ Model call successful")

            return DocumentAnalysis(filename=filename, analysis=resp)
            
        except Exception as e:
            print(f"\n❌ Parse error for {filename}:")
            print(f"Error type: {type(e)}")
            print(f"Error message: {str(e)}")
            import traceback
            print("Full traceback:")
            traceback.print_exc()
            
            if attempt < MAX_RETRIES:
                wait_time = RETRY_DELAY * (2 ** attempt)
                print(f"\033[33m⚠️  Retrying in {wait_time}s...\033[0m")
                await asyncio.sleep(wait_time)
            else:
                error_msg = f"Failed after {MAX_RETRIES} retries: {e}"
                return DocumentAnalysis(filename=filename, analysis=None, error=error_msg)

async def process_document_batch(client, batch: List[Tuple[str, str]]) -> List[DocumentAnalysis]:
    """Process a batch of documents concurrently."""
    tasks = []
    for filename, text in batch:
        task = analyze_text_with_instructor(client, text, filename)
        tasks.append(task)
    
    results = await asyncio.gather(*tasks, return_exceptions=False)
    return results

async def main_async():
    start_time = time.time()
    print("\033[1m🚀 Starting document analysis...\033[0m")
    
    # Create folders if they don't exist
    Path(INPUT_FOLDER).mkdir(exist_ok=True)
    Path(OUTPUT_FOLDER).mkdir(exist_ok=True)
    
    # Also create subfolders for High, Medium, Low relevance
    high_folder = os.path.join(OUTPUT_FOLDER, "high_relevance")
    medium_folder = os.path.join(OUTPUT_FOLDER, "medium_relevance")
    low_folder = os.path.join(OUTPUT_FOLDER, "low_relevance")
    Path(high_folder).mkdir(exist_ok=True)
    Path(medium_folder).mkdir(exist_ok=True)
    Path(low_folder).mkdir(exist_ok=True)

    # Initialize documents for each relevance level
    high_doc = Document()
    medium_doc = Document()
    low_doc = Document()
    
    # Set up the OpenAI client
    print("\n🔍 Debug: Setting up OpenAI client...")
    try:
        openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

        # Then wrap it with instructor
        client = instructor.from_openai(openai_client)

        print("✓ OpenAI instructor client created successfully")
        
    except Exception as e:
        print(f"\n❌ Error during client setup: {str(e)}")
        print(f"Error type: {type(e)}")
        import traceback
        print("Full traceback:")
        traceback.print_exc()
        return

    print("\033[32m✓ Successfully initialized OpenAI instructor client\033[0m")

    combined_doc = Document()
    
    # Collect all documents to process
    documents_to_process = []
    print("\n📁 Scanning input folder...")
    for filename in os.listdir(INPUT_FOLDER):
        if filename.endswith('.docx'):
            file_path = os.path.join(INPUT_FOLDER, filename)
            try:
                text = extract_text_from_docx(file_path)
                documents_to_process.append((filename, text))
                print(f"\033[32m✓ Successfully read {filename}\033[0m")
            except Exception as e:
                print(f"\033[31m❌ Error reading {filename}: {e}\033[0m")
    
    if not documents_to_process:
        print("\n\033[33m⚠️  No documents found to process\033[0m")
        return
    
    print(f"\n📊 Found {len(documents_to_process)} documents to process")
    print(f"🔄 Processing in batches of {MAX_CONCURRENT_REQUESTS}")
    
    successful_analyses = []
    failed_analyses = []
    
    for i in range(0, len(documents_to_process), MAX_CONCURRENT_REQUESTS):
        batch = documents_to_process[i:i + MAX_CONCURRENT_REQUESTS]
        batch_num = i // MAX_CONCURRENT_REQUESTS + 1
        total_batches = (len(documents_to_process) + MAX_CONCURRENT_REQUESTS - 1) // MAX_CONCURRENT_REQUESTS
        
        print(f"\n Processing batch {batch_num}/{total_batches} ({len(batch)} documents)...")
        batch_start_time = time.time()
        
        try:
            batch_results = await process_document_batch(client, batch)
            
            for result in batch_results:
                if result.error or result.analysis is None:
                    failed_analyses.append(result)
                else:
                    successful_analyses.append(result)
            
            batch_time = time.time() - batch_start_time
            print(f"\033[32m✓ Completed batch {batch_num}/{total_batches} in {batch_time:.1f}s\033[0m")
            
        except Exception as e:
            print(f"\033[31m❌ Error processing batch {batch_num}: {e}\033[0m")
    
    # Modify the document creation section
    print("\n Creating final documents...")
    high_cases = [a for a in successful_analyses if a.analysis.relevance_level == "High"]
    medium_cases = [a for a in successful_analyses if a.analysis.relevance_level == "Medium"]
    low_cases = [a for a in successful_analyses if a.analysis.relevance_level == "Low"]
    
    # Process each relevance level and store a single doc for each
    for cases, doc, level, folder_path in [
        (high_cases, high_doc, "high", high_folder),
        (medium_cases, medium_doc, "medium", medium_folder),
        (low_cases, low_doc, "low", low_folder)
    ]:
        for i, analysis in enumerate(cases):
            create_formatted_docx(doc, analysis.filename, analysis.analysis, is_first=(i == 0))
        
        if cases:  # Only save if there are cases of this relevance level
            output_file = os.path.join(folder_path, f"{level}_relevance_analysis.docx")
            doc.save(output_file)
            print(f"\033[32m✓ {level.capitalize()} relevance analysis saved to {output_file}\033[0m")
    
    # NEW CODE TO OUTPUT 9 DOCUMENTS (IF THEY EXIST) INTO THE THREE FOLDERS
    # We create a doc for each combination of Relevance x Support Level
    docs_map = {}
    relevance_levels = ["High", "Medium", "Low"]
    support_levels = ["Strongly Supports", "Supports", "Does not Support"]
    for r in relevance_levels:
        for s in support_levels:
            docs_map[(r, s)] = Document()
    
    results_by_combo = defaultdict(list)
    for analysis_result in successful_analyses:
        if analysis_result.analysis:
            r = analysis_result.analysis.relevance_level
            s = analysis_result.analysis.support_level
            if r in relevance_levels and s in support_levels:
                results_by_combo[(r, s)].append(analysis_result)
    
    for (r, s), doc_combo in docs_map.items():
        combo_cases = results_by_combo.get((r, s), [])
        if combo_cases:
            for i, analysis in enumerate(combo_cases):
                create_formatted_docx(doc_combo, analysis.filename, analysis.analysis, is_first=(i == 0))
            
            # Choose subfolder based on r
            if r == "High":
                subfolder = high_folder
            elif r == "Medium":
                subfolder = medium_folder
            else:
                subfolder = low_folder
            
            output_file = os.path.join(subfolder, f"Relevance {r} - {s}.docx")
            doc_combo.save(output_file)
            print(f"\033[32m✓ Relevance {r} - {s} analysis saved to {output_file}\033[0m")

    total_time = time.time() - start_time
    print("\n🔥 Summary:")
    print(f"Total documents processed: {len(documents_to_process)}")
    print(f"Successful analyses: {len(successful_analyses)}")
    print(f"Failed analyses: {len(failed_analyses)}")
    print(f"Total time: {total_time:.1f}s")
    
    if failed_analyses:
        print("\n❌ Failed documents:")
        for failed in failed_analyses:
            print(f"- {failed.filename}: {failed.error}")
    
    print("\n\033[1m✨ Processing complete!\033[0m")

def main():
    try:
        asyncio.run(main_async())
    except KeyboardInterrupt:
        print("\n\033[33m⚠️  Process interrupted by user\033[0m")
    except Exception as e:
        print(f"\n\033[31m❌ Fatal error: {str(e)}\033[0m")

if __name__ == "__main__":
    main()
